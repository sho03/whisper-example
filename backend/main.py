from openai import OpenAI

client = OpenAI()

from docx import Document

def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    return transcription.text

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "あなたは、言語理解と要約に訓練された高度に熟練したAIです。以下のテキストを読み、その内容を簡潔な抄録段落にまとめてください。最も重要なポイントを保持し、一貫した読みやすい要約を提供してください。この要約は、テキスト全体を読む必要なく、議論の主要なポイントを理解するのに役立つようにしてください。不要な詳細や余談は避けてください。"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "あなたは、言語理解と要約に訓練された高度に熟練したAIです。以下のテキストを読み、それを簡潔な抄録段落にまとめてください。最も重要なポイントを保持し、テキスト全体を読む必要なく、議論の主要なポイントを理解できるような、一貫した読みやすい要約を提供してください。不要な詳細や余談は避けてください。"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "あなたは、会話の分析とアクションアイテムの抽出におけるAIの専門家です。以下のテキストを確認し、合意されたタスク、割り当てられた仕事、または実行が必要とされたアクションを特定してください。これらは特定の個人に割り当てられたタスクであるか、グループが決定した一般的なアクションであるかもしれません。これらのアクションアイテムを明確かつ簡潔にリストしてください。"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "あなたは言語と感情分析の専門知識を持つAIです。以下のテキストのセンチメント（感情）を分析するタスクがあります。議論の全体的なトーン、使用されている言語によって伝えられる感情、および単語やフレーズが使用される文脈を考慮してください。センチメントが一般的にポジティブ、ネガティブ、またはニュートラルであるかを示し、可能な場合は分析のための簡単な説明を提供してください。"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

# save_as_docx(minutes, 'meeting_minutes.docx')

def main(file_path):
    transcription = transcribe_audio(file_path)
    minutes = meeting_minutes(transcription)

    return minutes

